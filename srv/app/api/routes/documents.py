from __future__ import annotations

import asyncio
import io
import json
import logging
import uuid
from datetime import datetime, timezone

from dependency_injector.wiring import Provide, inject
from fastapi import APIRouter, Depends, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from langchain_core.documents import Document
from pydantic import BaseModel

from app.api.deps import get_current_user_id
from app.containers import ApplicationContainer
from app.services.document_store import DocumentRecord, DocumentStore
from app.services.progress import ProgressEvent, ProgressNotifier
from pipeline.runner import PipelineRunner

logger = logging.getLogger(__name__)

router = APIRouter()

ALLOWED_EXTENSIONS = {".md", ".txt", ".docx"}
ALLOWED_MIME_TYPES = {
    "text/markdown",
    "text/plain",
    "application/octet-stream",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def _sse_frame(event: str, data: dict) -> str:
    return f"event: {event}\ndata: {json.dumps(data)}\n\n"


def _parse_file_content(filename: str, raw: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".docx"):
        import docx

        doc = docx.Document(io.BytesIO(raw))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return raw.decode("utf-8")


class ExtractRequest(BaseModel):
    selected_characters: list[str]
    pipeline_option: str = "advanced"


# ── Upload ────────────────────────────────────────────────────────────────────


@router.post("/upload")
@inject
async def upload_document(
    file: UploadFile,
    user_id: str = Depends(get_current_user_id),
    store: DocumentStore = Depends(Provide[ApplicationContainer.document_store]),
) -> dict:
    if not file.filename:
        raise HTTPException(status_code=400, detail="Filename is required")

    ext = "." + file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {ext}. "
            f"Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
        )

    raw = await file.read()
    content = _parse_file_content(file.filename, raw)

    record = DocumentRecord(
        id=str(uuid.uuid4()),
        user_id=user_id,
        filename=file.filename,
        mime_type=file.content_type or "application/octet-stream",
        size=len(raw),
        content=content,
        status="uploaded",
        uploaded_at=datetime.now(timezone.utc).isoformat(),
    )
    await store.add(record)

    return {
        "id": record.id,
        "filename": record.filename,
        "size": record.size,
        "status": record.status,
        "uploaded_at": record.uploaded_at,
    }


# ── List ──────────────────────────────────────────────────────────────────────


@router.get("")
@inject
async def list_documents(
    user_id: str = Depends(get_current_user_id),
    store: DocumentStore = Depends(Provide[ApplicationContainer.document_store]),
) -> list[dict]:
    docs = await store.list(user_id)
    return [
        {
            "id": d.id,
            "filename": d.filename,
            "size": d.size,
            "status": d.status,
            "uploaded_at": d.uploaded_at,
            "chunks_stored": d.chunks_stored,
            "error_message": d.error_message,
        }
        for d in docs
    ]


# ── Extract (SSE) ────────────────────────────────────────────────────────────


@router.post("/{document_id}/extract")
@inject
async def extract_knowledge(
    document_id: str,
    body: ExtractRequest,
    user_id: str = Depends(get_current_user_id),
    store: DocumentStore = Depends(Provide[ApplicationContainer.document_store]),
    notifier: ProgressNotifier = Depends(Provide[ApplicationContainer.progress_notifier]),
    runner: PipelineRunner = Depends(Provide[ApplicationContainer.ingestion_runner]),
) -> StreamingResponse:
    doc = await store.get(user_id, document_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="Document not found")
    if doc.status not in ("uploaded", "error"):
        raise HTTPException(
            status_code=409,
            detail=f"Document is already in status '{doc.status}', cannot extract",
        )

    await store.update(user_id, document_id, status="extracting")

    async def _stream():
        async for event in notifier.subscribe(document_id):
            if event.stage == "complete":
                yield _sse_frame("complete", {"chunks_stored": event.chunks_total or 0})
            elif event.stage == "failed":
                yield _sse_frame("error", {"message": event.message})
            else:
                yield _sse_frame("progress", event.model_dump())

    async def _run_pipeline():
        try:
            lc_doc = Document(
                page_content=doc.content,
                metadata={"source": doc.filename},
            )

            async def _on_progress(
                stage: str,
                progress_pct: int,
                chunks_total: int | None,
                chunks_processed: int | None,
            ) -> None:
                await notifier.notify(
                    document_id,
                    ProgressEvent(
                        stage=stage,
                        progress_pct=progress_pct,
                        chunks_total=chunks_total,
                        chunks_processed=chunks_processed,
                    ),
                )

            chunk_count = await runner.run(
                documents=[lc_doc],
                known_characters=body.selected_characters,
                pipeline_option=body.pipeline_option,
                on_progress=_on_progress,
            )
            await store.update(
                user_id, document_id,
                status="extracted",
                chunks_stored=chunk_count,
            )
        except Exception:
            logger.exception("Pipeline failed for document %s", document_id)
            await store.update(
                user_id, document_id, status="error", error_message="Pipeline failed",
            )
            await notifier.notify(
                document_id,
                ProgressEvent(stage="failed", progress_pct=0, message="Pipeline failed"),
            )

    asyncio.create_task(_run_pipeline())

    return StreamingResponse(
        _stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


# ── Delete ────────────────────────────────────────────────────────────────────


@router.delete("/{document_id}", status_code=204)
@inject
async def delete_document(
    document_id: str,
    user_id: str = Depends(get_current_user_id),
    store: DocumentStore = Depends(Provide[ApplicationContainer.document_store]),
):
    deleted = await store.delete(user_id, document_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="Document not found")
    return None
